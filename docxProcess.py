"""
Time：2021/10/13 16:38
Author: Qi Chengwen(BUAA)
Version: V1.0
File: WordDocProcess.py
IDE:Jetbrains PyCharm
Description: 用于word文件的创建搜索和处理
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.shared import Inches
import time
import os

# root = "C:\\Users\\huawei\\iCloudDrive\\EngineeringRecords"
# doc=Document("C:\\Users\\huawei\\iCloudDrive\\EngineeringRecords\\工程记录管理测试.docx")
search_result = []
count_num = [0, 0, 0]


def createEngineeringDoc(title, author, contacts, tags, content, filepath, filename):
    document = Document()
    document.add_heading(title, 0)
    authorParagraph = document.add_paragraph('Author：')  # 添加作者信息
    authorRun = authorParagraph.add_run(author)
    authorRun.font.name = u'楷体'
    authorr = authorRun._element
    authorr.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    today = time.strftime("%Y-%m-%d", time.localtime())
    document.add_paragraph('Time：' + today)  # 添加时间信息
    document.add_paragraph('Contacts：' + contacts)  # 添加联系方式信息
    tagParagraph = document.add_paragraph('TAGs：')  # 添加Tag
    tagRun = tagParagraph.add_run(tags)
    tagRun.font.name = u'楷体'
    tagrr = tagRun._element
    tagrr.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
    document.add_heading(time.strftime("%Y-%m-%d---%H:%M:%S", time.localtime()), 1)  # 添加当前时间
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Inches(0)
    run = paragraph.add_run(content)
    run.font.name = u'宋体'
    run.font.size = Pt(12)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # # 增加分页
    # document.add_page_break()
    # 保存文件
    if not os.path.exists(filepath):
        os.makedirs(filepath)
    document.save(filepath + '\\' + filename + '.docx')
    return 1.0


'''
搜索算法其实不需要太讲究，精确查找需要完全匹配，模糊查询需要起码按顺序出现，具体实现应该根据不同的搜索类型类判断
该函数用来搜索word文件，需要在另外一个函数中递归调用搜索文件夹

'''


def search_in_str(content, target, fuzzy):
    content_len = len(content)
    target_len = len(target)
    flag = 0
    if fuzzy == 1:
        j = 0
        for i in range(content_len):
            if content[i] == target[j]:
                j += 1
                if j == target_len:
                    flag = 1
                    break
    else:
        for i in range(content_len - target_len):
            if target[0] == content[i]:
                flag = 1
                for j in range(1, target_len):
                    if target[j] != content[i + j]:
                        flag = 0
                        break
    return flag


def search_word_content(type, target, fuzzy, path):  # type给定搜索类型，0为按标题查找，1为按作者查找， 2为按tag查找，3为按内容查找
    file = Document(path)
    flag = 0
    if type == 0:  # 说明按标题查找
        headings = file.paragraphs[0].text
        flag = search_in_str(headings, target, fuzzy)
    elif type == 1:
        author = file.paragraphs[1].text.replace("Author：", "")
        flag = search_in_str(author, target, fuzzy)
    elif type == 2:
        tags = file.paragraphs[4].text.replace("TAGs：", "")
        flag = search_in_str(tags, target, fuzzy)
    elif type == 3:
        for i in range(0, len(file.paragraphs)):  # 逐段搜索内容，一旦找到直接返回
            para_contents = file.paragraphs[i].text
            if para_contents == "":
                continue
            flag = search_in_str(para_contents, target, fuzzy)
            if flag == 1:
                break
    return flag

    # # 输出每一段的内容
    # for para in file.paragraphs:
    #     print(para.text)
    #
    # # 输出段落编号及段落内容
    # for i in range(len(file.paragraphs)):
    #     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)


def search_whole_db(type, target, fuzzy, docPath):
    paths = os.listdir(docPath)
    for item in paths:
        path = os.path.join(docPath, item)
        if os.path.isdir(path):
            search_whole_db(type, target, fuzzy, path)
        else:
            if path[-5:] == ".docx":
                flag = 0
                if "~$" in path:
                    continue
                flag = search_word_content(type, target, fuzzy, path)
                if flag == 1:
                    search_result.append(path)


def return_search_result(type, target, fuzzy, docPath):
    search_whole_db(type, target, fuzzy, docPath)
    # f = open('C:\\Users\\huawei\\Desktop\\log.txt', 'a+')
    # for i in search_result:
    #     f.write(i + '\n')
    # f.close()
    result2return = search_result.copy()
    search_result.clear()
    return result2return


def count_file_and_num(count_path):  # 统计文件总数和工程数量
    paths = os.listdir(count_path)
    if count_path[-4:] == "工程项目":
        count_num[1] = len(paths)
    for item in paths:
        path = os.path.join(count_path, item)
        if os.path.isdir(path):
            count_file_and_num(path)
        else:
            if path[-5:] == ".docx":
                if "~$" in path:
                    continue
                count_num[0] += 1
                count_num[2] += os.path.getsize(path)

def return_count_result(count_path):
    count_file_and_num(count_path)
    num2return = count_num.copy()
    num2return[2] /= (1024 * 1024)
    count_num[0] = 0
    count_num[1] = 0
    count_num[2] = 0
    return [str(num2return[0]), str(num2return[1]), str(num2return[2])]

def add_content(content, subtitle, filepath):
    document = Document(filepath)
    document.add_heading(time.strftime("%Y-%m-%d---%H:%M:%S", time.localtime()) + '-->' + subtitle, 1)  # 添加当前时间
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Inches(0)
    run = paragraph.add_run(content)
    run.font.name = u'宋体'
    run.font.size = Pt(12)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.save(filepath)
    return 1

# return_count_result("C:\\Users\\huawei\\iCloudDrive\\EngineeringRecords")
# x = return_search_result(2, "Python", 1, "C:\\Users\\huawei\\iCloudDrive\\EngineeringRecords")
# print(x)
# createEngineeringDoc('测试功能', '齐承文-BUAA', 'Chengwen_qi@outlook.com', '测试', '本文件用于测试创建的功能','D:\\testCreate\\testMultiDir', 'test')

# 参考https://www.jianshu.com/p/1f60cdd9655a
# add_content("测试新加内容的功能", "C:\\Users\\huawei\\iCloudDrive\\EngineeringRecords\\程序功能测试\\测试追加内容功能.docx")